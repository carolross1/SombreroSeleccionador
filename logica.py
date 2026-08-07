# -*- coding: utf-8 -*-
"""
logica.py
---------
Lógica de negocio del aplicativo "Test de Personalidad: ¿Qué Casa te representa?"

Contiene:
    - El mapeo de cada una de las 16 preguntas del cuestionario a las 4 casas
      (Gryffindor, Slytherin, Ravenclaw, Hufflepuff), reconstruido a partir del
      contenido de cada opción (no de su posición, a diferencia del script de
      Google Apps Script original, que asumía que el orden de las opciones
      siempre era Gryffindor-Slytherin-Ravenclaw-Hufflepuff).
    - Los mismos pesos por pregunta definidos en el script original.
    - Funciones de estadística descriptiva implementadas "a mano" (sin depender
      únicamente de pandas.describe()), tal como pide el SABER HACER.
    - Funciones de agrupación (K-Means) y reducción de dimensionalidad (PCA).

Este módulo es independiente de Streamlit para que pueda probarse por separado.
"""

from __future__ import annotations

import io
import json
import unicodedata
from dataclasses import dataclass
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.metrics import silhouette_score
from sklearn.preprocessing import StandardScaler

CASAS = ["Gryffindor", "Slytherin", "Ravenclaw", "Hufflepuff"]

COLORES = {
    "Gryffindor": "#740001",
    "Slytherin": "#1A472A",
    "Ravenclaw": "#0E1A40",
    "Hufflepuff": "#916300",
}

COLORES_SECUNDARIOS = {
    "Gryffindor": "#D3A625",
    "Slytherin": "#AAAAAA",
    "Ravenclaw": "#946B2D",
    "Hufflepuff": "#372E29",
}

ANIMAL = {
    "Gryffindor": "León",
    "Slytherin": "Serpiente",
    "Ravenclaw": "Águila",
    "Hufflepuff": "Tejón",
}

# Rasgo y descripción breve, texto original (no son citas de libros/películas),
# pensados solo para dar contexto visual al resultado del clustering.
RASGO = {
    "Gryffindor": "Valentía",
    "Slytherin": "Ambición",
    "Ravenclaw": "Sabiduría",
    "Hufflepuff": "Lealtad",
}

DESCRIPCION = {
    "Gryffindor": "Actúa primero y piensa después; no rehúye el riesgo cuando hay algo importante en juego.",
    "Slytherin": "Piensa en estrategia y resultados; sabe reconocer y aprovechar una oportunidad.",
    "Ravenclaw": "Prefiere entender antes que actuar; disfruta analizar un problema desde todos los ángulos.",
    "Hufflepuff": "Prioriza al grupo por encima de sí mismo; es constante y de confianza.",
}

# Lemas breves, texto original escrito para este proyecto (no son citas de
# los libros ni de las películas), pensados solo como cierre visual de cada
# tarjeta de casa.
LEMAS = {
    "Gryffindor": "El valor no es la ausencia de miedo, sino actuar a pesar de él.",
    "Slytherin": "Toda gran meta empieza con una jugada inteligente.",
    "Ravenclaw": "La mente que pregunta es la que llega más lejos.",
    "Hufflepuff": "Ningún logro vale la pena si se deja a alguien atrás.",
}

# Pesos por pregunta, en el orden en que aparecen las 16 preguntas
# (idénticos a los del script de Google Apps Script proporcionado).
PESOS: List[int] = [2, 2, 3, 1, 1, 2, 1, 2, 3, 3, 2, 3, 2, 1, 3, 1]
MAX_POSIBLE = sum(PESOS)  # 30

# Mapeo pregunta -> {texto de opción: casa}. Reconstruido leyendo el
# contenido/intención de cada opción del cuestionario (valentía -> Gryffindor,
# ambición/conveniencia propia -> Slytherin, sabiduría/análisis -> Ravenclaw,
# lealtad/cuidado del grupo -> Hufflepuff).
PREGUNTAS: List[Dict[str, str]] = [
    {
        "Me aseguro de que la persona afectada esté bien, ante todo": "Hufflepuff",
        "Confronto la situación de inmediato, sin pensarlo mucho": "Gryffindor",
        "Evalúo primero si me conviene involucrarme": "Slytherin",
        "Busco la manera más inteligente y estratégica de resolverlo": "Ravenclaw",
    },
    {
        "Planear la estrategia más eficiente para destacar": "Slytherin",
        "Investigar a fondo antes de proponer una solución": "Ravenclaw",
        "Asegurarme de que todos colaboren y se sientan incluidos": "Hufflepuff",
        "Tomar el liderazgo y avanzar rápido, aunque haya riesgos": "Gryffindor",
    },
    {
        "Que sea valiente y esté dispuesto a defenderme": "Gryffindor",
        "Que sea inteligente y tenga conversaciones interesantes": "Ravenclaw",
        "Que sea astuto y me ayude a conseguir lo que quiero": "Slytherin",
        "Que sea leal y esté ahí siempre, pase lo que pase": "Hufflepuff",
    },
    {
        "Buscar la forma más rápida de sacar ventaja de la situación": "Slytherin",
        "Analizar todas las opciones antes de decidir": "Ravenclaw",
        "Actuar de inmediato, aunque no tenga todo resuelto": "Gryffindor",
        "Buscar ayuda y resolverlo en conjunto con otros": "Hufflepuff",
    },
    {
        "Me interesa más entender cómo funciona el juego que ganar": "Ravenclaw",
        "Quiero ganar por mi propio esfuerzo y valentía": "Gryffindor",
        "Disfruto competir, pero me importa más que todos se diviertan": "Hufflepuff",
        "Quiero ganar, y usaré cualquier ventaja disponible": "Slytherin",
    },
    {
        "Analizo qué salió mal para no repetir el error y ganar la próxima vez": "Slytherin",
        "Me levanto rápido e intento de nuevo con más fuerza": "Gryffindor",
        "Busco apoyo en otros y no me rindo aunque tarde más": "Hufflepuff",
        "Reflexiono profundamente sobre las causas del fracaso": "Ravenclaw",
    },
    {
        "Una actividad grupal donde todos participen": "Hufflepuff",
        "Un juego de estrategia donde pueda ganar": "Slytherin",
        "Un deporte de alto riesgo o adrenalina": "Gryffindor",
        "Resolver acertijos o aprender algo nuevo": "Ravenclaw",
    },
    {
        "Considerar cómo afecta mi decisión a los demás": "Hufflepuff",
        "Detenerme a analizar la lógica de la situación, aunque tarde más": "Ravenclaw",
        "Confiar en mi instinto y actuar sin dudar": "Gryffindor",
        "Pensar qué me conviene más a largo plazo": "Slytherin",
    },
    {
        "Uno que cuida del bienestar de todo su equipo": "Hufflepuff",
        "Uno ambicioso que sabe cómo conseguir resultados": "Slytherin",
        "Uno que toma decisiones basadas en datos y lógica": "Ravenclaw",
        "Uno que inspira con el ejemplo y va al frente": "Gryffindor",
    },
    {
        "Que sea justa para todos los involucrados": "Hufflepuff",
        "Que tenga sentido lógico y esté bien fundamentada": "Ravenclaw",
        "Hacer lo correcto, aunque sea arriesgado": "Gryffindor",
        "Que me beneficie a mí y a mis metas": "Slytherin",
    },
    {
        "El que se asegura de que nadie quede atrás": "Hufflepuff",
        "El que actúa primero y pide perdón después": "Gryffindor",
        "El que mantiene la calma y piensa con claridad": "Ravenclaw",
        "El que ve la oportunidad detrás del caos": "Slytherin",
    },
    {
        "Mi ambición y capacidad de lograr lo que me propongo": "Slytherin",
        "Mi inteligencia y curiosidad": "Ravenclaw",
        "Mi valentía": "Gryffindor",
        "Mi lealtad y buen corazón": "Hufflepuff",
    },
    {
        "Analizando el problema desde una perspectiva objetiva": "Ravenclaw",
        "Buscando un acuerdo donde todos queden bien": "Hufflepuff",
        "Buscando la manera de salir ganando yo": "Slytherin",
        "Enfrentándolo directamente, sin rodeos": "Gryffindor",
    },
    {
        "El reconocimiento y el éxito": "Slytherin",
        "El desafío en sí mismo": "Gryffindor",
        "El aprender algo nuevo en el proceso": "Ravenclaw",
        "Ayudar o apoyar a alguien más": "Hufflepuff",
    },
    {
        "Capacidad de generar confianza en los demás": "Hufflepuff",
        "Coraje inquebrantable": "Gryffindor",
        "Habilidad para conseguir lo que quiero": "Slytherin",
        "Inteligencia excepcional": "Ravenclaw",
    },
    {
        "Me lanzo sin pensarlo mucho, con energía": "Gryffindor",
        "Investigo bien antes de empezar": "Ravenclaw",
        "Ya tengo un plan de cómo sacarle el mayor provecho": "Slytherin",
        "Pienso en cómo involucrar y apoyar a los demás": "Hufflepuff",
    },
]

# Imagenes y descripciones largas de cada casa, usadas tanto en la pestaña
# "Datos" como en "Seleccion individual" (antes estaban duplicadas en app.py;
# se centralizan aqui para evitar que las dos copias se desincronicen).
CASA_IMG_MAP: Dict[str, str] = {
    "Gryffindor": "gryffindor.png",
    "Hufflepuff": "hufflepuff.png",
    "Ravenclaw": "ravenclaw.png",
    "Slytherin": "slytherin.png",
}

CASA_DESCRIPCIONES: Dict[str, Dict[str, str]] = {
    "Gryffindor": {
        "subtitulo": "VALENTIA · LEON",
        "descripcion": "ACTUA PRIMERO Y PIENSA DESPUES; NO REHUYE EL RIESGO CUANDO HAY ALGO IMPORTANTE EN JUEGO.",
        "cita": '"EL VALOR NO ES LA AUSENCIA DE MIEDO, SINO ACTUAR A PESAR DE EL."',
    },
    "Slytherin": {
        "subtitulo": "AMBICION · SERPIENTE",
        "descripcion": "PIENSA EN ESTRATEGIA Y RESULTADOS; SABE RECONOCER Y APROVECHAR UNA OPORTUNIDAD.",
        "cita": '"TODA GRAN META EMPIEZA CON UNA JUGADA INTELIGENTE."',
    },
    "Ravenclaw": {
        "subtitulo": "SABIDURIA · AGUILA",
        "descripcion": "PREFIERE ENTENDER ANTES QUE ACTUAR; DISFRUTA ANALIZAR UN PROBLEMA DESDE TODOS LOS ANGULOS.",
        "cita": '"LA MENTE QUE PREGUNTA ES LA QUE LLEGA MAS LEJOS."',
    },
    "Hufflepuff": {
        "subtitulo": "LEALTAD · TEJON",
        "descripcion": "PRIORIZA AL GRUPO POR ENCIMA DE SI MISMO; ES CONSTANTE Y DE CONFIANZA.",
        "cita": '"NINGUN LOGRO VALE LA PENA SI SE DEJA A ALGUIEN ATRAS."',
    },
}

N_PREGUNTAS = len(PREGUNTAS)  # 16


# ---------------------------------------------------------------------------
# Cuestionario configurable
# ---------------------------------------------------------------------------
# Todo lo anterior (CASAS, PREGUNTAS, PESOS) es el cuestionario de Hogwarts,
# que se usa como valor por defecto. Para poder analizar OTRO cuestionario
# (distintas preguntas, opciones o hasta distinto numero de categorias) sin
# tocar codigo, ese mismo contenido se empaqueta en un objeto `Cuestionario`
# que se puede reemplazar cargando un archivo .json con el mismo formato
# (ver `plantilla_cuestionario_json` para generar un ejemplo editable).

@dataclass
class Cuestionario:
    nombre: str
    categorias: List[str]
    preguntas: List[Dict[str, str]]  # por pregunta: {texto de opcion: categoria}
    pesos: List[int]
    enunciados: List[str]  # enunciado real de cada pregunta (puede ir vacio)

    def __post_init__(self) -> None:
        n = len(self.preguntas)
        if len(self.pesos) != n:
            raise ValueError(
                f"El cuestionario '{self.nombre}' tiene {n} preguntas pero {len(self.pesos)} pesos; deben coincidir."
            )
        if not self.enunciados:
            self.enunciados = [f"Pregunta {i + 1}" for i in range(n)]
        elif len(self.enunciados) != n:
            raise ValueError(
                f"El cuestionario '{self.nombre}' tiene {n} preguntas pero {len(self.enunciados)} enunciados; deben coincidir."
            )
        categorias_usadas = {cat for opciones in self.preguntas for cat in opciones.values()}
        desconocidas = categorias_usadas - set(self.categorias)
        if desconocidas:
            raise ValueError(
                f"El cuestionario '{self.nombre}' usa categorias no declaradas en 'categorias': {sorted(desconocidas)}."
            )

    @property
    def n_preguntas(self) -> int:
        return len(self.preguntas)

    @property
    def max_posible(self) -> int:
        return sum(self.pesos)


CUESTIONARIO_HOGWARTS = Cuestionario(
    nombre="Test de Personalidad: Que Casa te representa",
    categorias=CASAS,
    preguntas=PREGUNTAS,
    pesos=PESOS,
    enunciados=[],  # no se conserva el enunciado original del formulario; usa "Pregunta N"
)


def dict_a_cuestionario(contenido: dict) -> Cuestionario:
    """Construye un Cuestionario a partir de un diccionario con el formato:

        {
          "nombre": "Mi cuestionario",
          "categorias": ["A", "B", "C"],
          "preguntas": [
            {
              "enunciado": "Texto real de la pregunta 1 (opcional)",
              "peso": 2,
              "opciones": {"texto opcion 1": "A", "texto opcion 2": "B", ...}
            },
            ...
          ]
        }

    Lanza ValueError con un mensaje claro si el formato no es valido, para
    poder mostrarlo directo en la interfaz sin traceback. La usan tanto el
    importador de archivos .json como el editor visual dentro de la app.
    """
    if not isinstance(contenido, dict):
        raise ValueError("El cuestionario debe ser un objeto en su nivel superior.")

    categorias = contenido.get("categorias")
    preguntas_raw = contenido.get("preguntas")
    if not isinstance(categorias, list) or len(categorias) < 2:
        raise ValueError("El cuestionario necesita al menos 2 categorias (por ejemplo, 2 grupos de personalidad).")
    if not isinstance(preguntas_raw, list) or len(preguntas_raw) < 1:
        raise ValueError("El cuestionario necesita al menos 1 pregunta.")

    preguntas: List[Dict[str, str]] = []
    pesos: List[int] = []
    enunciados: List[str] = []
    for i, p in enumerate(preguntas_raw):
        if not isinstance(p, dict) or "opciones" not in p:
            raise ValueError(f"La pregunta {i + 1} no tiene respuestas configuradas.")
        opciones = p["opciones"]
        if not isinstance(opciones, dict) or len(opciones) < 2:
            raise ValueError(f"La pregunta {i + 1} necesita al menos 2 opciones de respuesta.")
        preguntas.append(opciones)
        try:
            pesos.append(int(p.get("peso", 1)))
        except (TypeError, ValueError):
            raise ValueError(f"El peso (importancia) de la pregunta {i + 1} debe ser un numero entero.")
        enunciados.append(str(p.get("enunciado") or f"Pregunta {i + 1}"))

    return Cuestionario(
        nombre=str(contenido.get("nombre") or "Cuestionario personalizado"),
        categorias=[str(c) for c in categorias],
        preguntas=preguntas,
        pesos=pesos,
        enunciados=enunciados,
    )


def cargar_cuestionario_json(datos: str | bytes) -> Cuestionario:
    """Construye un Cuestionario a partir del contenido de un archivo .json.
    Uso pensado para quienes prefieren preparar el archivo fuera de la app
    (por ejemplo, para compartirlo con otra persona)."""
    try:
        contenido = json.loads(datos)
    except json.JSONDecodeError as e:
        raise ValueError(f"El archivo no es un JSON valido: {e}") from e
    return dict_a_cuestionario(contenido)


def cuestionario_a_dict(cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS) -> dict:
    """Convierte un cuestionario a un diccionario simple (nombre, categorias,
    preguntas con enunciado/peso/opciones), la misma forma que usan tanto el
    editor visual como el archivo .json de respaldo."""
    return {
        "nombre": cuestionario.nombre,
        "categorias": list(cuestionario.categorias),
        "preguntas": [
            {
                "enunciado": cuestionario.enunciados[i],
                "peso": cuestionario.pesos[i],
                "opciones": dict(cuestionario.preguntas[i]),
            }
            for i in range(cuestionario.n_preguntas)
        ],
    }


def plantilla_cuestionario_json(cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS) -> str:
    """Exporta un cuestionario (por defecto, el de Hogwarts) como JSON, para
    quienes prefieran guardarlo o compartirlo como archivo de respaldo."""
    return json.dumps(cuestionario_a_dict(cuestionario), ensure_ascii=False, indent=2)

# Palabras clave para detectar automáticamente las columnas de metadatos
# dentro de un export de Google Forms (los nombres exactos de columna pueden
# variar ligeramente, por eso se busca por coincidencia parcial).
METADATA_KEYWORDS = {
    "marca_temporal": ["marca temporal", "timestamp"],
    "correo": ["correo"],
    "edad": ["edad"],
    "ocupacion": ["ocupaci"],
    "genero": ["gener"],
    "sueno": ["sueño", "sueno"],
    "estres": ["estrés", "estres"],
}


def _normalizar(texto: str) -> str:
    """Quita acentos, espacios extra y pasa a minúsculas para comparar texto
    de forma robusta ante pequeñas diferencias de captura."""
    if not isinstance(texto, str):
        return ""
    texto = texto.strip().lower()
    texto = "".join(
        c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn"
    )
    return texto


@dataclass
class ColumnasDetectadas:
    metadatos: Dict[str, str]
    preguntas: List[str]


def detectar_columnas(df: pd.DataFrame) -> ColumnasDetectadas:
    """Detecta cuáles columnas del CSV son metadatos (correo, edad, género,
    etc.) y cuáles corresponden a las 16 preguntas del cuestionario, asumiendo
    que las columnas de preguntas conservan el orden original del formulario.
    """
    metadatos: Dict[str, str] = {}
    columnas_restantes: List[str] = []

    for col in df.columns:
        col_norm = _normalizar(col)
        asignada = False
        for clave, palabras in METADATA_KEYWORDS.items():
            if clave in metadatos:
                continue
            if any(p in col_norm for p in palabras):
                metadatos[clave] = col
                asignada = True
                break
        if not asignada:
            columnas_restantes.append(col)

    return ColumnasDetectadas(metadatos=metadatos, preguntas=columnas_restantes)


def validar_columnas_preguntas(
    columnas_preguntas: List[str], cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS
) -> Tuple[bool, str]:
    """Valida que el CSV traiga las columnas de preguntas esperadas por el
    cuestionario activo, en su orden original (el calculo de puntajes es
    posicional). Devuelve (ok, mensaje). ok=False cuando faltan demasiadas
    columnas como para confiar en el resultado.
    """
    n = len(columnas_preguntas)
    n_esperado = cuestionario.n_preguntas
    if n == 0:
        return False, (
            "No pudimos encontrar las preguntas dentro de tu archivo. Revisa que sea el CSV "
            "correcto (el que exportaste del formulario) y que no le hayas quitado los encabezados."
        )
    if n < n_esperado:
        return False, (
            f"Tu archivo tiene menos preguntas de las que espera el cuestionario '{cuestionario.nombre}' "
            f"({n} de {n_esperado}). Con preguntas faltantes, los resultados no serian confiables. "
            "Revisa que sea el CSV correcto, o ajusta el cuestionario en la seccion 'Cuestionario' "
            "de la barra lateral para que coincida con tus preguntas."
        )
    if n > n_esperado:
        return True, (
            f"Encontramos {n} columnas que parecen preguntas, pero el cuestionario '{cuestionario.nombre}' "
            f"solo espera {n_esperado}. Vamos a usar unicamente las primeras {n_esperado} en el orden en "
            "que aparecen; si el resultado se ve raro, revisa que no se haya colado alguna columna extra."
        )
    return True, ""


def contar_respuestas_faltantes(
    df: pd.DataFrame, columnas_preguntas: List[str], cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS
) -> pd.Series:
    """Para cada fila, cuenta cuantas de las preguntas del cuestionario
    activo quedaron sin respuesta valida (vacia o con un texto que no
    matchea ninguna opcion conocida de esa pregunta). Util para avisar de
    datos incompletos antes de entrenar/aplicar el modelo."""
    n = min(len(columnas_preguntas), cuestionario.n_preguntas)
    faltantes = []
    for _, fila in df.iterrows():
        vacias = 0
        for i in range(n):
            col = columnas_preguntas[i]
            respuesta = _normalizar(fila.get(col, ""))
            mapeo = {_normalizar(k) for k in cuestionario.preguntas[i].keys()}
            if not respuesta or respuesta not in mapeo:
                vacias += 1
        faltantes.append(vacias)
    return pd.Series(faltantes, index=df.index, name="respuestas_faltantes")


def validar_modelo_features(
    modelo: dict, X: np.ndarray, categorias: List[str] | None = None
) -> Tuple[bool, str]:
    """Verifica que un modelo cargado desde .pkl sea compatible con las
    features (columnas) que se le quieren pasar, antes de llamar a
    transform()/predict() y tronar con un error críptico de sklearn. Si se
    pasa `categorias` (las del cuestionario activo), tambien valida que los
    NOMBRES coincidan exactamente y no solo la cantidad -- dos cuestionarios
    distintos podrian tener el mismo numero de categorias pero significar
    cosas distintas."""
    columnas_esperadas = modelo.get("features")
    if not columnas_esperadas:
        return False, "El archivo .pkl no tiene el formato esperado (falta la clave 'features')."
    if len(columnas_esperadas) != X.shape[1]:
        return False, (
            f"El modelo cargado espera {len(columnas_esperadas)} variable(s) "
            f"({', '.join(columnas_esperadas)}) y este CSV/cuestionario genero {X.shape[1]}. "
            "Es probable que el .pkl venga de otro cuestionario."
        )
    if categorias is not None:
        columnas_actuales = [f"{c}_pct" for c in categorias]
        if columnas_actuales != list(columnas_esperadas):
            return False, (
                f"El modelo cargado se entreno con las categorias {columnas_esperadas} y el "
                f"cuestionario activo tiene {columnas_actuales}. Aunque coincidan en cantidad, "
                "no son las mismas categorias: carga el cuestionario .json original junto con este modelo."
            )
    escalador = modelo.get("escalador")
    if escalador is not None and getattr(escalador, "n_features_in_", None) not in (None, X.shape[1]):
        return False, "El escalador del modelo cargado no coincide con el numero de variables de este CSV."
    return True, ""


def calcular_puntajes(
    df: pd.DataFrame, columnas_preguntas: List[str], cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS
) -> pd.DataFrame:
    """Calcula, para cada respondiente (fila), el puntaje absoluto y el
    porcentaje obtenido en cada categoria del cuestionario activo (Hogwarts
    por defecto, o el que se haya cargado desde .json), replicando la logica
    de ponderacion basada en el contenido de la respuesta (no en su posicion
    dentro de las opciones).
    """
    n = min(len(columnas_preguntas), cuestionario.n_preguntas)
    categorias = cuestionario.categorias
    filas = []
    for _, fila in df.iterrows():
        puntos = {c: 0 for c in categorias}
        for i in range(n):
            col = columnas_preguntas[i]
            respuesta = _normalizar(fila.get(col, ""))
            mapeo = {_normalizar(k): v for k, v in cuestionario.preguntas[i].items()}
            categoria = mapeo.get(respuesta)
            if categoria:
                puntos[categoria] += cuestionario.pesos[i]
        filas.append(puntos)

    puntajes = pd.DataFrame(filas, index=df.index)
    for c in categorias:
        puntajes[f"{c}_pct"] = round(puntajes[c] / cuestionario.max_posible * 100, 1)
    puntajes["Casa_dominante"] = puntajes[categorias].idxmax(axis=1)
    return puntajes


# ---------------------------------------------------------------------------
# Estadística descriptiva "propia" (sin usar df.describe())
# ---------------------------------------------------------------------------

def media_propia(valores: List[float]) -> float:
    valores = [v for v in valores if v is not None and not pd.isna(v)]
    if not valores:
        return float("nan")
    return sum(valores) / len(valores)


def desviacion_estandar_propia(valores: List[float]) -> float:
    valores = [v for v in valores if v is not None and not pd.isna(v)]
    n = len(valores)
    if n < 2:
        return float("nan")
    m = media_propia(valores)
    varianza = sum((v - m) ** 2 for v in valores) / (n - 1)
    return varianza ** 0.5


def moda_propia(valores: List) -> List:
    """Devuelve TODOS los valores empatados en la frecuencia maxima
    (soporta resultados unimodales, bimodales, trimodales, etc.)."""
    valores = [v for v in valores if v is not None and not (isinstance(v, float) and pd.isna(v))]
    if not valores:
        return []
    conteo: Dict[object, int] = {}
    for v in valores:
        conteo[v] = conteo.get(v, 0) + 1
    frecuencia_max = max(conteo.values())
    return [valor for valor, freq in conteo.items() if freq == frecuencia_max]


def tabla_frecuencias(valores: List) -> pd.Series:
    conteo: Dict[object, int] = {}
    for v in valores:
        if v is None or (isinstance(v, float) and pd.isna(v)):
            continue
        conteo[v] = conteo.get(v, 0) + 1
    return pd.Series(conteo).sort_values(ascending=False)


def resumen_estadistico(df: pd.DataFrame, columnas_numericas: List[str]) -> pd.DataFrame:
    filas = []
    for col in columnas_numericas:
        valores = pd.to_numeric(df[col], errors="coerce").tolist()
        valores_validos = [v for v in valores if not pd.isna(v)]
        filas.append(
            {
                "Variable": col,
                "n": len(valores_validos),
                "Media": round(media_propia(valores_validos), 2) if valores_validos else float("nan"),
                "Desv. estándar": round(desviacion_estandar_propia(valores_validos), 2) if valores_validos else float("nan"),
                "Mínimo": min(valores_validos) if valores_validos else float("nan"),
                "Máximo": max(valores_validos) if valores_validos else float("nan"),
            }
        )
    return pd.DataFrame(filas)


# ---------------------------------------------------------------------------
# Agrupación (K-Means) y reducción de dimensionalidad (PCA)
# ---------------------------------------------------------------------------

def preparar_features(puntajes: pd.DataFrame, categorias: List[str] = CASAS) -> np.ndarray:
    columnas = [f"{c}_pct" for c in categorias]
    return puntajes[columnas].to_numpy(dtype=float)


def calcular_inercias(X_esc: np.ndarray, k_min: int = 2, k_max: int = 8) -> pd.DataFrame:
    """Inercia (método del codo) para distintos valores de k, útil para
    justificar la elección del número de clusters en el reporte."""
    filas = []
    k_max = min(k_max, max(k_min, X_esc.shape[0] - 1))
    for k in range(k_min, k_max + 1):
        modelo = KMeans(n_clusters=k, random_state=42, n_init=10)
        modelo.fit(X_esc)
        filas.append({"k": k, "Inercia": modelo.inertia_})
    return pd.DataFrame(filas)


def calcular_siluetas(X_esc: np.ndarray, k_min: int = 2, k_max: int = 8) -> pd.DataFrame:
    """Coeficiente de silueta para distintos valores de k, como complemento
    al metodo del codo (la inercia siempre baja al subir k y no basta por si
    sola para justificar la eleccion; la silueta si penaliza clusters mal
    separados)."""
    filas = []
    k_max = min(k_max, max(k_min, X_esc.shape[0] - 1))
    for k in range(k_min, k_max + 1):
        modelo = KMeans(n_clusters=k, random_state=42, n_init=10)
        etiquetas = modelo.fit_predict(X_esc)
        try:
            score = silhouette_score(X_esc, etiquetas)
        except ValueError:
            score = float("nan")
        filas.append({"k": k, "Silueta": score})
    return pd.DataFrame(filas)


def entrenar_kmeans(X: np.ndarray, k: int) -> Tuple[StandardScaler, KMeans, PCA, np.ndarray, np.ndarray]:
    """Entrena el escalador, el modelo K-Means y un PCA de 2 componentes para
    visualización/reducción de dimensionalidad. Devuelve también las
    etiquetas de cluster y las coordenadas PCA."""
    escalador = StandardScaler()
    X_esc = escalador.fit_transform(X)

    kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
    etiquetas = kmeans.fit_predict(X_esc)

    pca = PCA(n_components=2, random_state=42)
    coords_pca = pca.fit_transform(X_esc)

    return escalador, kmeans, pca, etiquetas, coords_pca


def generar_reporte_docx(
    k: int,
    n_aspirantes: int,
    metricas: Dict[str, object],
    cruce: pd.DataFrame,
    perfil_cluster: pd.DataFrame,
    origen_modelo: str,
    cuestionario: Cuestionario = CUESTIONARIO_HOGWARTS,
) -> bytes:
    """Arma un reporte ejecutivo en Word (.docx) con el resumen de la corrida
    de K-Means + PCA: metricas de calidad del modelo, tabla cruzada
    categoria-dominante x cluster, y el perfil promedio de cada cluster.
    Devuelve los bytes del documento, listos para un st.download_button.
    """
    from docx import Document
    from docx.shared import Pt

    es_hogwarts = cuestionario is CUESTIONARIO_HOGWARTS
    etiqueta_cat = "casa" if es_hogwarts else "categoria"

    doc = Document()

    titulo = "El Sombrero Seleccionador — Reporte ejecutivo" if es_hogwarts else f"{cuestionario.nombre} — Reporte ejecutivo"
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(
        "Resumen de la corrida de aprendizaje no supervisado (StandardScaler → "
        f"K-Means → PCA) sobre los porcentajes por {etiqueta_cat} de los respondientes."
    )

    doc.add_heading("Configuracion del modelo", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Numero de clusters (k): ").bold = True
    p.add_run(str(k))
    p = doc.add_paragraph()
    p.add_run("Respondientes analizados: ").bold = True
    p.add_run(str(n_aspirantes))
    p = doc.add_paragraph()
    p.add_run("Cuestionario: ").bold = True
    p.add_run(f"{cuestionario.nombre} ({len(cuestionario.categorias)} categorias, {cuestionario.n_preguntas} preguntas)")
    p = doc.add_paragraph()
    p.add_run("Origen del modelo: ").bold = True
    p.add_run("cargado desde un .pkl previamente entrenado" if origen_modelo == "cargado" else "entrenado en esta sesion")

    doc.add_heading("Metricas de calidad", level=2)
    tabla_m = doc.add_table(rows=1, cols=2)
    tabla_m.style = "Light Grid Accent 1"
    hdr = tabla_m.rows[0].cells
    hdr[0].text, hdr[1].text = "Metrica", "Valor"
    filas_metricas = []
    if metricas.get("silueta") is not None:
        filas_metricas.append(("Coeficiente de silueta", f"{metricas['silueta']:.3f}"))
    if metricas.get("inercia") is not None:
        filas_metricas.append(("Inercia (K-Means)", f"{metricas['inercia']:.1f}"))
    varianza = metricas.get("varianza_explicada_pca")
    if varianza:
        filas_metricas.append(("Varianza explicada por el PCA 2D", f"{sum(varianza) * 100:.1f}%"))
    for nombre, valor in filas_metricas:
        fila = tabla_m.add_row().cells
        fila[0].text, fila[1].text = nombre, valor

    doc.add_paragraph(
        "La varianza explicada indica que tan fiel es la proyeccion 2D del PCA respecto a "
        f"las {len(cuestionario.categorias)} variables originales (porcentaje por {etiqueta_cat}). "
        "El coeficiente de silueta va de -1 a 1 y penaliza clusters mal separados; la inercia por "
        "si sola no basta para justificar el numero de clusters, por eso se complementa con la silueta."
    )

    doc.add_heading(f"{etiqueta_cat.capitalize()} dominante vs. cluster asignado", level=2)
    doc.add_paragraph(
        f"Si cada cluster concentra mayormente una {etiqueta_cat} dominante, el algoritmo esta "
        "agrupando perfiles similares sin haber usado esa etiqueta durante el entrenamiento."
    )
    tabla_c = doc.add_table(rows=1, cols=len(cruce.columns) + 1)
    tabla_c.style = "Light Grid Accent 1"
    hdr = tabla_c.rows[0].cells
    hdr[0].text = f"{etiqueta_cat.capitalize()} dominante"
    for j, col in enumerate(cruce.columns):
        hdr[j + 1].text = f"Cluster {col}"
    for idx, fila_datos in cruce.iterrows():
        fila = tabla_c.add_row().cells
        fila[0].text = str(idx)
        for j, val in enumerate(fila_datos):
            fila[j + 1].text = str(val)

    doc.add_heading(f"Perfil promedio por cluster (% por {etiqueta_cat})", level=2)
    tabla_p = doc.add_table(rows=1, cols=len(perfil_cluster.columns) + 1)
    tabla_p.style = "Light Grid Accent 1"
    hdr = tabla_p.rows[0].cells
    hdr[0].text = "Cluster"
    for j, col in enumerate(perfil_cluster.columns):
        hdr[j + 1].text = col
    for idx, fila_datos in perfil_cluster.iterrows():
        fila = tabla_p.add_row().cells
        fila[0].text = str(idx)
        for j, val in enumerate(fila_datos):
            fila[j + 1].text = f"{val:.1f}"

    for tabla in (tabla_m, tabla_c, tabla_p):
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def aplicar_modelo_guardado(modelo: dict, X: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
    """Aplica un modelo previamente entrenado (escalador + KMeans + PCA, tal
    como los guarda la app en el .pkl de la pestaña Descargas) sobre nuevas
    filas, SIN reentrenar nada: solo transform()/predict().

    `modelo` es el diccionario cargado con joblib.load(...) (claves
    "escalador", "kmeans", "pca", "k", "features").
    Devuelve (etiquetas, coords_pca), en el mismo formato que entrenar_kmeans.
    """
    escalador: StandardScaler = modelo["escalador"]
    kmeans: KMeans = modelo["kmeans"]
    pca: PCA = modelo["pca"]

    X_esc = escalador.transform(X)
    etiquetas = kmeans.predict(X_esc)
    coords_pca = pca.transform(X_esc)

    return etiquetas, coords_pca
